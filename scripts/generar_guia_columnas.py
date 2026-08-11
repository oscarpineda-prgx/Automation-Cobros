"""Genera una guía en Word con el significado y el cálculo de cada columna clave.

Documento de referencia para el auditor: qué columna es llave, cuáles se copian del CFDI
(CPA Vision), cuáles se toman del propio Compras, cuáles se calculan, y cómo se calculan
las columnas de auditoría.

Uso:
    python scripts/generar_guia_columnas.py
Deja `Guia_Columnas_Compras.docx` en la raíz del proyecto.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SALIDA = Path(__file__).resolve().parent.parent / "Guia_Columnas_Compras.docx"

VERDE = RGBColor(0x00, 0xB0, 0x50)
AMARILLO = RGBColor(0xBF, 0x90, 0x00)
NARANJA = RGBColor(0xC5, 0x50, 0x00)
AZUL = RGBColor(0x1F, 0x4E, 0x78)


def _titulo(doc, texto, color=AZUL, size=15):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _parrafo_9(doc, texto, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(9)
    r.italic = italic
    return p


def _tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(encabezados):
        celda = t.rows[0].cells[i]
        celda.text = ""
        run = celda.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for fila in filas:
        cells = t.add_row().cells
        for i, valor in enumerate(fila):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(valor))
            run.font.size = Pt(9)
    return t


def main() -> None:
    doc = Document()

    h = doc.add_heading("Guía de columnas — Compras / Cruce CPA Vision", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Automation Costos · PRGX · Soriana Audit Suite")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.add_run(
        "Soriana pagó a cada proveedor según el costo de su sistema (SAP). El cruce trae lo "
        "que el proveedor realmente facturó en sus CFDI (facturas electrónicas, descargadas "
        "de CPA Vision) y llena las columnas EDI vacías del Compras. Luego el recálculo "
        "compara ambos costos y calcula la diferencia a reclamar."
    ).font.size = Pt(10)

    # Leyenda de origen
    _titulo(doc, "Origen de cada columna (código de color)", size=13)
    leyenda = doc.add_paragraph()
    for texto, color in (
        ("● Copiada del CFDI (CPA Vision)   ", AMARILLO),
        ("● Del propio Compras   ", VERDE),
        ("● Calculada   ", NARANJA),
    ):
        r = leyenda.add_run(texto)
        r.bold = True
        r.font.color.rgb = color
        r.font.size = Pt(10)
    doc.add_paragraph()

    # 1. Llave del cruce
    _titulo(doc, "1. Llave del cruce (cómo se emparejan Compras y CPA Vision)", size=13)
    _tabla(
        doc,
        ["Concepto", "En Compras", "En CPA Vision", "Nota"],
        [
            ["Proveedor", "cnpj (RFC)", "RFC (partición)", "El RFC se detecta solo, desde el propio Compras"],
            ["Código de barras", "codbarra", "noIdentificacion", "Se leen solo dígitos, sin ceros iniciales; NO se usa la columna upc"],
            ["Número de factura", "invnbr", "Serie + Folio", "invnbr trae la serie pegada y con formato irregular; se normaliza"],
        ],
    )
    doc.add_paragraph()

    _parrafo_9(
        doc,
        "El cruce se hace en DOS PASADAS, porque invnbr viene capturado de forma "
        "inconsistente (conviven FN-21226, FN21226 y -21226 en el mismo proveedor):",
    )
    _tabla(
        doc,
        ["Pasada", "Llave", "Cómo se normaliza", "Ejemplo"],
        [
            [
                "1ª — principal",
                "código de barras + Serie+Folio",
                "Mayúsculas, se quita todo lo que no sea letra o número (la serie SE CONSERVA)",
                "FN-21226 → FN21226",
            ],
            [
                "2ª — respaldo",
                "código de barras + Folio",
                "Solo los dígitos, sin ceros a la izquierda. Se aplica ÚNICAMENTE a los renglones que no cruzaron en la 1ª",
                "FN-21226 → 21226",
            ],
        ],
    )
    _parrafo_9(
        doc,
        "El respaldo es más laxo (ignora la serie), por eso va en segundo lugar y solo sobre "
        "lo que quedó sin cruzar. El riesgo de falso positivo queda acotado porque la búsqueda "
        "ya va restringida por RFC y por código de barras.",
        italic=True,
    )
    doc.add_paragraph()

    # 1b. Reglas de seguridad del cruce
    _titulo(doc, "2. Dos reglas de seguridad del cruce", size=13)
    _tabla(
        doc,
        ["Regla", "Qué significa"],
        [
            [
                "Solo rellena celdas VACÍAS",
                "Si el Compras ya traía un dato en una columna EDI, el cruce NO lo pisa. Lo que el auditor capturó a mano se respeta.",
            ],
            [
                "Los CFDI en conflicto se DESCARTAN",
                "Si una misma llave apunta a varios conceptos del CFDI con valores DISTINTOS, no se elige ninguno: se deja vacío y se reporta en el conteo 'CFDI en conflicto'. Preferimos no llenar a llenar mal. (Las líneas repetidas con valores idénticos sí se colapsan en una: dan igual.)",
            ],
        ],
    )
    doc.add_paragraph()

    # 3. Copiadas del CFDI
    _titulo(doc, "3. Columnas que se COPIAN del CFDI (CPA Vision)", AMARILLO, size=13)
    doc.add_paragraph(
        "Se copian tal cual, solo en las celdas que venían vacías en Compras.",
    ).runs[0].font.size = Pt(9)
    _tabla(
        doc,
        ["Columna en Compras", "Viene de (CPA Vision)", "Qué es"],
        [
            ["canfac_edi", "Cantidad", "Cantidad facturada por el proveedor"],
            ["ctonto_edi", "Valor Unitario", "Costo unitario que el proveedor facturó"],
            ["poriva_edi", "TASA IVA", "Porcentaje de IVA (0.16, 0.08, 0...)"],
            ["impiva_edi", "IVA", "Importe de IVA del renglón (en pesos)"],
            ["prieps_edi", "TASA IEPS", "Porcentaje de IEPS"],
            ["imieps_edi", "IEPS", "Importe de IEPS del renglón (en pesos)"],
            ["totfactura", "Total", "Total de TODA la factura (se repite en cada renglón)"],
            ["uuid", "UUID", "Folio fiscal único del CFDI (SAT)"],
        ],
    )
    doc.add_paragraph()

    # 4. Del propio Compras
    _titulo(doc, "4. Columna que se toma del propio Compras", VERDE, size=13)
    _tabla(
        doc,
        ["Columna", "Se toma de", "Qué es"],
        [["factem_edi", "fact_empaq", "Factor de empaque: unidades por caja. Ya viene en Compras, NO del CFDI"]],
    )
    doc.add_paragraph()

    # 5. EDI calculadas
    _titulo(doc, "5. Columnas EDI que se CALCULAN", NARANJA, size=13)
    _tabla(
        doc,
        ["Columna", "Fórmula", "Ejemplo"],
        [
            ["ctobto_edi", "ctonto_edi × factem_edi", "11 × 20 = 220 (costo por caja)"],
            ["impart_edi", "ctobto_edi × canfac_edi × (1 + poriva_edi)", "220 × 320 × 1 = 70,400"],
        ],
    )
    _parrafo_9(
        doc,
        "Solo se calculan donde la celda estaba vacía Y hay insumo (ctonto_edi con dato): "
        "así no se inventan ceros en renglones que no cruzaron.",
        italic=True,
    )
    doc.add_paragraph()

    # 6. Columnas de control
    _titulo(doc, "6. Columnas de CONTROL (doble validación, no son del entregable)", size=13)
    _parrafo_9(
        doc,
        "Los importes de impuesto se COPIAN del CFDI, no se despejan del total. Se midió "
        "sobre 66 millones de renglones: las columnas IVA e IEPS de CPA Vision coinciden al "
        "100% con el importe del impuesto, mientras que despejarlos desde el Total solo "
        "acierta en el 0.9%. La causa es que la mayoría de las facturas de Soriana mezclan "
        "artículos gravados y a tasa 0 (alimentos), y dividir el Total completo entre 1.16 "
        "asume que todo está gravado. La fórmula se conserva únicamente como control:",
    )
    _tabla(
        doc,
        ["Columna de control", "Fórmula", "Para qué sirve"],
        [
            [
                "impiva_edi_formula",
                "totfactura ÷ (1 + poriva_edi) × poriva_edi",
                "Se compara contra impiva_edi (el valor bueno, copiado del CFDI). Si difieren más de 2 centavos se reporta.",
            ],
            [
                "imieps_edi_formula",
                "totfactura ÷ (1 + prieps_edi) × prieps_edi",
                "Igual, contra imieps_edi. Es normal que difieran: confirma que copiar era lo correcto.",
            ],
        ],
    )
    doc.add_paragraph()

    # 7. Auditoría
    _titulo(doc, "7. Columnas de AUDITORÍA (el corazón del cálculo)", NARANJA, size=13)
    _parrafo_9(
        doc,
        "Aquí se compara lo que Soriana pagó (ctouni, su sistema) contra lo que el proveedor "
        "facturó (ctonto_edi, del CFDI). Regla conservadora: solo corrige a la baja.",
    )
    _parrafo_9(
        doc,
        "IMPORTANTE — qué significa 'cruzó con CPA': que el renglón TRAE dato del CFDI "
        "(uuid o ctonto_edi presentes), no que el valor sea distinto de cero. Un CFDI puede "
        "traer IVA o IEPS = 0 legítimamente (producto exento) y ese 0 debe respetarse. "
        "Confirmado con Mónica y Perla el 2026-07-31.",
        italic=True,
    )
    _tabla(
        doc,
        ["Columna", "Cálculo", "Significado"],
        [
            [
                "cto_aud",
                "Si cruzó CPA y ctonto_edi>0 y < ctouni → ctonto_edi; si no → ctouni",
                "Costo auditado: el menor entre lo pagado y lo facturado (nunca 0)",
            ],
            [
                "iva_aud",
                "Si cruzó CPA → poriva_edi (aunque sea 0); si no → iva_t007s (tasa SAP)",
                "Tasa de IVA que se aplica al cálculo",
            ],
            [
                "ieps_aud",
                "Si cruzó CPA → prieps_edi (aunque sea 0); si no → ieps_t007s (tasa SAP)",
                "Tasa de IEPS que se aplica",
            ],
            [
                "imp_aud",
                "cto_aud × can_rec × (1 + iva_aud) × (1 + ieps_aud)",
                "Importe auditado con impuestos, por renglón",
            ],
            [
                "debio_pagar_ne",
                "Suma de imp_aud por FOLIO (nota de entrada)",
                "Lo que Soriana DEBIÓ pagar por esa nota de entrada",
            ],
            [
                "tot_pagado_ne",
                "Máximo de paynetamt por FOLIO",
                "Lo que Soriana pagó por esa nota (máx, no suma: viene repetido)",
            ],
            [
                "dif_det_ne",
                "tot_pagado_ne − debio_pagar_ne",
                "Diferencia a nivel nota de entrada. Positiva = se pagó de más = a cobrar",
            ],
            [
                "debio_pagar_inv",
                "Suma de imp_aud por FACTURA (vndnbr|strnbr|invnbr)",
                "Lo que se debió pagar, agrupado por factura",
            ],
            [
                "tot_pagado_inv",
                "Máximo de paynetamt por FACTURA",
                "Lo pagado a nivel factura",
            ],
            [
                "dif_det_inv",
                "tot_pagado_inv − debio_pagar_inv",
                "Diferencia a nivel factura",
            ],
        ],
    )
    doc.add_paragraph()

    nota = doc.add_paragraph()
    nota.add_run("Nota sobre los dos niveles: ").bold = True
    nota.add_run(
        "hay dos cortes del mismo dato. 'NE' agrupa por nota de entrada (folio = tienda + "
        "nota); 'inv' agrupa por factura. El 'pagado' usa MÁXIMO y no suma porque el monto "
        "viene repetido en cada renglón del grupo."
    ).font.size = Pt(9)

    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.add_run(
        "Generado automáticamente desde scripts/generar_guia_columnas.py. "
        "Fuente: docs/MAPEO_CRUCE_CPA_COMPRAS.md, docs/LOGICA_NEGOCIO.md y "
        "automation_costos/calculations.py."
    )
    pie.runs[0].italic = True
    pie.runs[0].font.size = Pt(8)

    doc.save(SALIDA)
    print(f"Guía generada: {SALIDA}")


if __name__ == "__main__":
    main()
