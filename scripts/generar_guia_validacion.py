"""Genera una guía en Word que explica cómo se construye la Validación de Condiciones.

Documento de referencia para el auditor: qué hojas tiene el entregable, con qué criterio
se filtra, y de dónde sale cada columna del Consolidado y del Detalle.

Uso:
    python scripts/generar_guia_validacion.py
Deja `Guia_Validacion_Condiciones.docx` en la raíz del proyecto.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SALIDA = Path(__file__).resolve().parent.parent / "Guia_Validacion_Condiciones.docx"

AZUL = RGBColor(0x1F, 0x4E, 0x78)
VERDE = RGBColor(0x00, 0xB0, 0x50)
ROJO = RGBColor(0xC0, 0x00, 0x00)


def _titulo(doc, texto, color=AZUL, size=14):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _parrafo(doc, texto, size=10, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.size = Pt(size)
    r.italic = italic
    return p


def _tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(encabezados):
        t.rows[0].cells[i].text = ""
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    for fila in filas:
        cells = t.add_row().cells
        for i, valor in enumerate(fila):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(valor))
            run.font.size = Pt(9)
    return t


def main() -> None:
    doc = Document()

    h = doc.add_heading("Guía — Validación de Condiciones", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Automation Costos · PRGX · Soriana Audit Suite")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _parrafo(
        doc,
        "La Validación de Condiciones es el archivo que se ENTREGA al proveedor. Resume las "
        "diferencias de costo que la auditoría detectó (lo que Soriana pagó de más) para "
        "reclamarlas. Se genera a partir del Compras ya cruzado y recalculado.",
    )

    _titulo(doc, "Las hojas del archivo (3, o 4 si hay devoluciones)", size=13)
    _tabla(
        doc,
        ["Hoja", "Qué contiene", "¿Siempre?"],
        [
            ["Resumen", "El total de la diferencia a reclamar (una cifra)", "Sí"],
            ["Consolidado", "Una fila por nota de entrada con diferencia: cuánto se pagó, cuánto se debió, la diferencia", "Sí"],
            [
                "Ajustes",
                "Bitácora de las devoluciones de pago (MR8M / KG-14) que anulan o reducen diferencias",
                "Solo si el proveedor tiene devoluciones",
            ],
            ["Detalle PAGOS", "El desglose artículo por artículo de esas notas de entrada", "Sí"],
        ],
    )
    _parrafo(
        doc,
        "Si el Detalle supera el tope de filas de Excel (1,048,576) se continúa en hojas "
        "'Detalle PAGOS (2)', '(3)', etc. — nunca se trunca en silencio.",
        italic=True,
        size=9,
    )
    doc.add_paragraph()

    # Filtro
    _titulo(doc, "1. Qué renglones entran a la Validación (el filtro)", size=13)
    _parrafo(
        doc,
        "No entran todas las compras: solo las notas de entrada con diferencia. Hay dos casos:",
    )
    _tabla(
        doc,
        ["Situación", "Criterio que se aplica"],
        [
            [
                "El auditor YA clasificó la columna 'concepto'",
                "Entran solo las notas marcadas como 'dif costos'",
            ],
            [
                "Aún no se clasifica (preliminar)",
                "Entran las notas con dif_det_ne mayor a 1 peso (umbral configurable)",
            ],
            [
                "La nota tenía diferencia pero una devolución COMPENSADA la dejó en ~0",
                "SALE del Consolidado y queda registrada en la hoja 'Ajustes' (ver sección 4)",
            ],
        ],
    )
    doc.add_paragraph()

    # Resumen
    _titulo(doc, "2. Hoja RESUMEN", size=13)
    _tabla(
        doc,
        ["Celda", "Contenido", "De dónde sale"],
        [
            ["Título", "Tiendas Soriana / Proveedor / Validación de Condiciones", "Datos del proveedor"],
            [
                "Resultado Auditoría",
                "Suma de la columna de diferencia del Consolidado. Si el proveedor tiene devoluciones, suma 'Diferencia Ajustada' (ya neta); si no, 'Diferencia'",
                "Total a reclamar",
            ],
            ["Observaciones Auditor", "'Diferencia costos'", "Texto fijo (lo ajusta el auditor)"],
        ],
    )
    doc.add_paragraph()

    # Consolidado
    _titulo(doc, "3. Hoja CONSOLIDADO (una fila por nota de entrada)", size=13)
    _parrafo(doc, "Cada fila resume una nota de entrada con diferencia. Columnas:")
    _tabla(
        doc,
        ["Columna", "De dónde sale", "Qué es"],
        [
            ["Proveedor", "vndnbr", "Número de proveedor"],
            ["Folio", "tienda + nota de entrada", "Llave de la entrega física"],
            ["Tienda/CEDIS", "strnbr", "Tienda o centro de distribución"],
            ["Negocio", "po_org (o 'SORIANA')", "Formato de negocio"],
            ["Fecha Pedido", "podt", "Fecha de la orden de compra"],
            ["Pedido", "ponbr", "Orden de compra"],
            ["Fec Recibo", "rcvdt", "Fecha de recepción"],
            ["Nota Entrada", "rcvnbr", "Número de nota de entrada"],
            ["Factura", "invnbr_ne (o invnbr)", "Factura del proveedor"],
            ["Total Pagado", "tot_pagado_ne = MÁX de paynetamt por folio", "Lo que Soriana pagó por esa nota"],
            ["Debio Pagar", "debio_pagar_ne = SUMA de imp_aud por folio", "Lo que debió pagar según el costo auditado"],
            ["Diferencia", "dif_det_ne = Total Pagado − Debió Pagar", "Positiva = se pagó de más = A COBRAR"],
            ["Observaciones Auditor", "(vacío)", "Lo llena el auditor a mano"],
            ["Fecha Pago", "paychkdt_ne", "Fecha del pago"],
            ["Documento de Pago", "paychknbr_ne", "Referencia del pago"],
        ],
    )
    _parrafo(
        doc,
        "En la fila de encabezado, las columnas de importe llevan un SUBTOTAL(109,...) que "
        "suma toda la columna visible tras filtrar.",
        italic=True,
        size=9,
    )
    doc.add_paragraph()

    # Ajustes de pago
    _titulo(doc, "4. Devoluciones de pago: hoja AJUSTES y las 6 columnas extra", ROJO, size=13)
    _parrafo(
        doc,
        "El Line Item de Compras (Audit Tools) no contempla ciertas devoluciones de pago que "
        "sí anulan diferencias reales: si no se consideran, se le reclama al proveedor algo "
        "que Soriana ya le descontó. Estas devoluciones viven en otra fuente del sistema y "
        "son de dos tipos (acordado con Perla y Mónica el 2026-08-04):",
    )
    _tabla(
        doc,
        ["Tipo", "Qué es", "Cómo se liga al Consolidado"],
        [
            ["MR8M", "Devolución que NO referencia nota de entrada ni tienda, solo la factura", "Proveedor + Factura"],
            ["KG-14", "Ajuste que SÍ trae tienda y nota de entrada", "Nota de entrada + Tienda (= el Folio exacto)"],
        ],
    )
    doc.add_paragraph()

    _parrafo(doc, "Lo decisivo es si la devolución ya se ejecutó o no:")
    _tabla(
        doc,
        ["Estado", "Cómo se detecta", "Qué hace el sistema"],
        [
            [
                "COMPENSADA",
                "Tiene documento de pago (el cheque ya se generó; Soriana paga por corte semanal, los miércoles)",
                "SE RESTA de la diferencia del renglón, sin dejarla por debajo de cero. Si la factura toca varios folios, se reparte en cascada sin descontar de más",
            ],
            [
                "NO COMPENSADA",
                "El documento de pago viene vacío o en 0: el cliente ya lo identificó pero aún no lo ejecuta",
                "NO se resta. Se marca el renglón con la alerta 'No compensado/ejecutado' + el conteo y el monto pendiente, para que el auditor lo vigile y haga la resta manual cuando se compense",
            ],
        ],
    )
    _parrafo(
        doc,
        "Esa distinción es la razón de ser de la alerta: restar algo que todavía no se pagó "
        "subestimaría la reclamación (acordado el 2026-08-04).",
        italic=True,
        size=9,
    )
    doc.add_paragraph()

    _parrafo(doc, "Cuando hay devoluciones, el Consolidado pasa de 15 a 21 columnas. Las 6 nuevas:")
    _tabla(
        doc,
        ["Columna nueva", "Qué es"],
        [
            ["Tipo Ajuste", "MR8M, KG, o 'KG+MR8M' si al renglón le tocan los dos"],
            ["Ajuste Pagos", "Suma de las devoluciones COMPENSADAS aplicadas a ese renglón (va en negativo)"],
            ["Diferencia Ajustada", "Diferencia + Ajuste Pagos. Es la cifra que de verdad se reclama"],
            ["Compensado", "Bandera de alerta: 'No compensado/ejecutado' si hay pendientes"],
            ["Conteo No Compensados", "Cuántas devoluciones pendientes tocan ese renglón"],
            ["Monto No Compensado", "Cuánto suman esas pendientes (informativo, NO se resta)"],
        ],
    )
    doc.add_paragraph()

    _parrafo(doc, "La hoja 'Ajustes' es la bitácora: una fila por devolución cruzada. Columnas:")
    _tabla(
        doc,
        ["Columna", "Qué es"],
        [
            ["Proveedor / Tienda-CEDIS / Nota Entrada / Factura", "Identifican el renglón del Consolidado afectado"],
            ["Diferencia Original", "La diferencia ANTES de aplicar la devolución"],
            ["Tipo Ajuste", "MR8M o KG"],
            ["Compensado", "Sí / No"],
            ["Ajuste Aplicado", "Cuánto se restó (negativo). Es 0 si no está compensada"],
            ["Monto No Compensado", "Cuánto quedó pendiente. Es 0 si sí está compensada"],
            ["Documento Pago / Fecha Pago", "Referencia del pago. Vacíos mientras no se compense"],
            ["DOC_TEXT", "Texto del documento en el sistema origen, para rastrearlo"],
        ],
    )
    _parrafo(
        doc,
        "Nota: si el proveedor no tiene devoluciones —o si no hay conexión a la base al "
        "generar el archivo— la Validación sale exactamente como antes, con sus 15 columnas y "
        "sin hoja 'Ajustes'. Nunca falla por esto.",
        italic=True,
        size=9,
    )
    doc.add_paragraph()

    # Detalle
    _titulo(doc, "5. Hoja DETALLE PAGOS (una fila por artículo)", size=13)
    _parrafo(
        doc,
        "Desglose de cada artículo de las notas de entrada del Consolidado. Aquí se ve "
        "producto por producto el costo pagado contra el costo correcto.",
    )
    _tabla(
        doc,
        ["Columna", "De dónde sale", "Qué es"],
        [
            ["Folio_Pedido", "ponbr", "Orden de compra"],
            ["FecPed", "podt", "Fecha de pedido"],
            ["Folio_NE", "rcvnbr", "Nota de entrada"],
            ["FecRecibo", "rcvdt", "Fecha de recepción"],
            ["Factura", "invnbr_ne (o invnbr)", "Factura del proveedor"],
            ["Tienda/CEDIS", "strnbr", "Tienda / CEDIS"],
            ["Division", "nombre_division", "División de mercancía"],
            ["Categoria", "po_groupdescrip", "Categoría"],
            ["GciaCateg", "grupoarticulo", "Gerencia de categoría"],
            ["Material", "cltstyle", "Clave de material"],
            ["CodBarra", "codbarra", "Código de barras del artículo"],
            ["Descripcion", "itmdesc", "Descripción del artículo"],
            ["FactEmpaqu", "fact_empaq", "Factor de empaque (piezas por caja)"],
            ["CantRec", "can_rec", "Cantidad recibida"],
            ["CtoUnitario_sistema", "ctouni", "Costo unitario del sistema (SAP)"],
            ["Costo Unitario correcto", "cto_aud", "Costo auditado (el menor entre pagado y facturado)"],
            ["Porc_IEPS", "prieps_edi", "% de IEPS facturado"],
            ["IEPS_Aud", "ieps_aud", "Tasa de IEPS auditada"],
            ["Porc_IVA", "poriva_edi", "% de IVA facturado"],
            ["IVA_Aud", "iva_aud", "Tasa de IVA auditada"],
            ["Debio Pagar correcto", "imp_aud", "Importe auditado con impuestos, por artículo"],
        ],
    )
    doc.add_paragraph()

    # Nota / flujo
    _titulo(doc, "6. El cálculo de la diferencia, en 5 pasos", ROJO, size=13)
    _tabla(
        doc,
        ["Paso", "Qué pasa"],
        [
            ["1", "Por cada artículo: cto_aud = el menor entre lo que Soriana pagó (ctouni) y lo que el proveedor facturó (ctonto_edi)"],
            ["2", "imp_aud = cto_aud × cantidad recibida × (1+IVA) × (1+IEPS) — lo que se DEBIÓ pagar por artículo"],
            ["3", "Se suma imp_aud por nota de entrada → 'Debió Pagar'. Se toma el pago real (máx) → 'Total Pagado'"],
            ["4", "Diferencia = Total Pagado − Debió Pagar. Si es positiva, Soriana pagó de más y se reclama"],
            ["5", "Se restan las devoluciones YA COMPENSADAS → 'Diferencia Ajustada'. Las pendientes no se restan: se marcan como alerta"],
        ],
    )
    doc.add_paragraph()

    nota = doc.add_paragraph()
    nota.add_run("✔ Confirmado con negocio (Mónica, 2026-07-31): ").bold = True
    nota.add_run(
        "en el Detalle, la columna 'CtoUnitario_sistema' se llena con ctouni (el costo del "
        "sistema/SAP), no con ctonto_edi. Antes salía en cero cuando no había EDI; ya quedó "
        "corregido."
    ).font.size = Pt(9)

    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.add_run(
        "Generado desde scripts/generar_guia_validacion.py. Fuente: "
        "automation_costos/validation_exporter.py y calculations.py."
    )
    pie.runs[0].italic = True
    pie.runs[0].font.size = Pt(8)

    doc.save(SALIDA)
    print(f"Guía generada: {SALIDA}")


if __name__ == "__main__":
    main()
